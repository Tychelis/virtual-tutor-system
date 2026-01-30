import os
from pdf2image import convert_from_path
import uuid
from pathlib import Path
from typing import List, Tuple, Dict, Iterable, Set
from dataclasses import dataclass
from collections import Counter
from difflib import SequenceMatcher
from docx import Document as DocxDocument
import fitz
import regex as re

try:
    from langchain_core.documents import Document
except ModuleNotFoundError:
    from langchain.schema import Document

TOP_PCT = 0.1
BOTTOM_PCT = 0.1
REPEAT_THRESHOLD = 0.8
MIN_PAGE_COUNT = 10
FUZZY_SIM_TH = 0.9
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
HEADER_FOOTER_PATTERNS = [
    r"^\s*\d+\s*$",
    r"^\s*\d+\s*/\s*\d+\s*$",
    r"^\s*COMP\d{4}.*$",
    r"^\s*Week\s*\d+.*$",
    r"^\s*UNSW.*$",
    r"^\s*Course Outline.*$",
]


def pdf_to_imgs(pdf_path,save_dir):
  base_name = os.path.splitext(os.path.basename(pdf_path))[0]
  save_dir = os.path.join(save_dir, base_name)
  os.makedirs(save_dir, exist_ok=True)
  images = convert_from_path(pdf_path)
  for i, image in enumerate(images):
      image_path = os.path.join(save_dir, f"page_{i+1}.png")
      if os.path.exists(image_path):
            print(f"Page {i+1} already exists, skipping.")
            continue
      image.save(image_path, "PNG")
  print(f"Images saved in {save_dir}")
  return save_dir

def normalize_line(line: str) -> str:
    return re.sub(r"\s+", " ", line.strip())

def fuzzy_eq(a: str, b: str, th: float = FUZZY_SIM_TH) -> bool:
    return SequenceMatcher(None, a, b).ratio() >= th

def match_any_regex(text: str, patterns: List[str]) -> bool:
    return any(re.match(p, text, flags=re.IGNORECASE) for p in patterns)

@dataclass
class LineWithPos:
    text: str
    y0: float
    y1: float

@dataclass
class RawPage:
    page_num: int
    height: float
    lines: List[LineWithPos]
    raw_text: str

def extract_page_lines(page: fitz.Page) -> List[LineWithPos]:
    out = []
    td = page.get_text("dict")
    for block in td.get("blocks", []):
        for l in block.get("lines", []):
            spans = l.get("spans", [])
            if not spans:
                continue
            txt = "".join(s["text"] for s in spans).strip("\n")
            if not txt.strip():
                continue
            y0 = min(s["bbox"][1] for s in spans)
            y1 = max(s["bbox"][3] for s in spans)
            out.append(LineWithPos(txt, y0, y1))
    return out

def detect_header_footer_with_position(pages: List[RawPage]) -> Tuple[Set[str], Set[str]]:
    total = len(pages)
    header_cands, footer_cands = [], []
    for p in pages:
        top_th = p.height * TOP_PCT
        bot_th = p.height * (1 - BOTTOM_PCT)
        for l in p.lines:
            txt = normalize_line(l.text)
            if not txt:
                continue
            if match_any_regex(txt, HEADER_FOOTER_PATTERNS):
                if l.y0 < top_th:
                    header_cands.append(txt)
                elif l.y1 > bot_th:
                    footer_cands.append(txt)
                continue
            if l.y0 < top_th:
                header_cands.append(txt)
            elif l.y1 > bot_th:
                footer_cands.append(txt)

    def choose(cands):
        if total <= MIN_PAGE_COUNT:
            return set()
        cnt = Counter(cands)
        raw = [t for t, c in cnt.items() if c / total >= REPEAT_THRESHOLD]
        merged = []
        for t in raw:
            if not any(fuzzy_eq(t, m) for m in merged):
                merged.append(t)
        return set(merged)

    return choose(header_cands), choose(footer_cands)

def strip_header_footer(page: RawPage, headers: Set[str], footers: Set[str]) -> str:
    kept = []
    top_th = page.height * TOP_PCT
    bot_th = page.height * (1 - BOTTOM_PCT)
    for l in page.lines:
        nl = normalize_line(l.text)
        if not nl:
            continue
        if match_any_regex(nl, HEADER_FOOTER_PATTERNS):
            continue
        if any(fuzzy_eq(nl, h) for h in headers) and l.y0 < top_th:
            continue
        if any(fuzzy_eq(nl, f) for f in footers) and l.y1 > bot_th:
            continue
        kept.append(l.text)
    return "\n".join(kept)

def normalize_bullets(text: str) -> str:
    text = re.sub(fr"[{BULLET_CHARS}]", "-", text)
    text = re.sub(r"(?m)^\s*[vV]\s+", "- ", text)
    text = re.sub(r"(?m)^\s*\+\s+", "- ", text)
    text = re.sub(r"(?m)^\s*\*\s+", "- ", text)

    text = re.sub(r"(?m)^\s*-{2,}\s*", "- ", text)
    return text

BULLET_CHARS = r"§•∙·◦▪▶►※"

def basic_text_clean(text: str) -> str:
    text = re.sub(r"[\x00-\x1F\x7F]", " ", text)
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = normalize_bullets(text)

    return text.strip()

def build_sentence_splitter():
    try:
        import spacy
        nlp = spacy.load("en_core_web_sm", disable=["tagger", "ner", "lemmatizer"])
        return lambda t: [s.text.strip() for s in nlp(t).sents if s.text.strip()]
    except Exception:
        pass
    try:
        import nltk
        try:
            nltk.data.find("tokenizers/punkt")
        except LookupError:
            nltk.download("punkt")
        from nltk.tokenize import sent_tokenize
        return lambda t: [s.strip() for s in sent_tokenize(t) if s.strip()]
    except Exception:
        pass
    return lambda t: [p.strip() for p in re.split(r"(?<=[.!?])\s+", t) if p.strip()]

def sentences_to_chunks(sentences: List[str],
                        chunk_size=CHUNK_SIZE,
                        chunk_overlap=CHUNK_OVERLAP) -> List[str]:
    chunks, cur = [], ""
    for s in sentences:
        if len(cur) + len(s) + 1 <= chunk_size:
            cur = f"{cur} {s}".strip()
        else:
            if cur:
                chunks.append(cur)
            if chunk_overlap > 0 and cur:
                tail = cur[-chunk_overlap:]
                cur = (tail + " " + s).strip()
            else:
                cur = s
    if cur:
        chunks.append(cur)
    return chunks


def get_chunks_from_pdf(pdf_path: str) -> List[Document]:

    sent_split = build_sentence_splitter()
    chunks_all: List[Document] = []

    pth = Path(pdf_path).resolve()
    doc = fitz.open(pth)

    raw_pages: List[RawPage] = []
    for i in range(len(doc)):
        page = doc[i]
        raw_pages.append(
            RawPage(
                page_num=i + 1,
                height=page.rect.height,
                lines=extract_page_lines(page),
                raw_text=page.get_text("text"),
            )
        )
    doc.close()

    headers, footers = detect_header_footer_with_position(raw_pages)

    for rp in raw_pages:
        cleaned = strip_header_footer(rp, headers, footers)
        cleaned = basic_text_clean(cleaned)
        if not cleaned.strip():
            continue
        sents = sent_split(cleaned)
        page_chunks = sentences_to_chunks(sents, CHUNK_SIZE, CHUNK_OVERLAP)

        for ch in page_chunks:
            meta = {"source": str(pth), "page": rp.page_num}
            chunks_all.append(Document(page_content=ch, metadata=meta))

    return chunks_all

def get_chunks_from_docx(docx_path: str):
    sent_split = build_sentence_splitter()
    chunks_all: List[Document] = []

    doc = DocxDocument(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    cleaned = basic_text_clean(text)
    if cleaned.strip():
        sents = sent_split(cleaned)
        chunks = sentences_to_chunks(sents, CHUNK_SIZE, CHUNK_OVERLAP)
        for ch in chunks:
            meta = {"source": str(docx_path), "page": 1}
            chunks_all.append(Document(page_content=ch, metadata=meta))

    return chunks_all

def get_chunks_from_txt(txt_path:str):
    sent_split = build_sentence_splitter()
    chunks_all: List[Document] = []

    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read()
    cleaned = basic_text_clean(text)
    if cleaned.strip():
        sents = sent_split(cleaned)
        chunks = sentences_to_chunks(sents, CHUNK_SIZE, CHUNK_OVERLAP)
        for ch in chunks:
            meta = {"source": str(txt_path), "page": 1}
            chunks_all.append(Document(page_content=ch, metadata=meta))

    return chunks_all

def get_chunks_from_documents(document_path):
    pth = Path(document_path).resolve()
    ext = pth.suffix.lower()

    if ext == ".pdf":
        return get_chunks_from_pdf(pth)
    elif ext == ".docx":
        return get_chunks_from_docx(pth)
    elif ext == ".txt":
        return get_chunks_from_txt(pth)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
# pdfs = "/home/jialu/workspace/jialu/materials/test_0727.pdf"
# chunks = get_chunks_from_pdf(pdfs)
# for chunk in chunks:
#     print(chunk.page_content)
# print(len(chunks), "chunks")
# print(chunks[0].metadata, chunks[0].page_content[:300])